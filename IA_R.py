from openai import OpenAI

client = OpenAI(
    # defaults to os.environ.get("OPENAI_API_KEY")
    # api_key="My API Key",
)
from docx import Document


    # Function that uses the Whisper model to take in the audio file and transcribe it
def transcribe_audio(audio_file_path):
    with open(audio_file_path, 'rb') as audio_file:
        transcription = client.audio.transcriptions.create("whisper-1", audio_file)
    return transcription['text']


    # Function which will serve as the main function of this application
def meeting_minutes(transcription):                 # "transcription" is the text we obtained from Whisper.
    abstract_summary = abstract_summary_extraction(transcription)
    key_points = key_points_extraction(transcription)
    action_items = action_item_extraction(transcription)
    sentiment = sentiment_analysis(transcription)
    return {                                        # Passed to the four other functions
        'abstract_summary': abstract_summary,
        'key_points': key_points,
        'action_items': action_items,
        'sentiment': sentiment
    }


    # Takes the transcription and summarizes it into a concise abstract paragraph with the aim 
    # to retain the most important points while avoiding unnecessary details or tangential points.
def abstract_summary_extraction(transcription):
    response = client.chat.completions.create(
        model="gpt-4",
        temperature=0,
        messages=[          # There are many different possible ways of achieving similar results 
                            # through the process commonly referred to as prompt engineering
            {
                "role": "system",
                "content": "You are a highly skilled AI trained in language comprehension and summarization. I would like you to read the following text and summarize it into a concise abstract paragraph. Aim to retain the most important points, providing a coherent and readable summary that could help a person understand the main points of the discussion without needing to read the entire text. Please avoid unnecessary details or tangential points."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return completion.choices[0].message.content


    # Identifies and lists the main points discussed in the meeting.
def key_points_extraction(transcription):
    response = client.chat.completions.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "We are a company that transcribe medical consultation into usables notes. You are a proficient AI with a specialty in distilling information into key points. Based on the following text, identify and list the main points that were discussed or brought up. These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. Your goal is to provide a list that someone could read to quickly understand what was talked about."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return completion.choices[0].message.content


    # Identifies tasks, assignments, or actions agreed upon or mentioned during the meeting
def action_item_extraction(transcription):
    response = client.chat.completions.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are an AI expert in analyzing conversations and extracting action items. Please review the text and identify any tasks, assignments, or actions that were agreed upon or mentioned as needing to be done. These could be tasks assigned to specific individuals, or general actions that the group has decided to take. Please list these action items clearly and concisely."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return completion.choices[0].message.content


    # Analyzes the overall sentiment of the discussion. It considers the tone, 
    # the emotions conveyed by the language used, and the context in which words and phrases are used
def sentiment_analysis(transcription):
    response = client.chat.completions.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "As an AI with expertise in language and emotion analysis, your task is to analyze the sentiment of the following text. Please consider the overall tone of the discussion, the emotion conveyed by the language used, and the context in which words and phrases are used. Indicate whether the sentiment is generally positive, negative, or neutral, and provide brief explanations for your analysis where possible."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return completion.choices[0].message.content


    # Converts the raw text to a Word document
def save_as_docx(minutes, filename):
    doc = Document()
    for key, value in minutes.items():
        # Replace underscores with spaces and capitalize each word for the heading
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        # Add a line break between sections
        doc.add_paragraph()
    doc.save(filename)


    # This code will transcribe the audio file Earningscall.wav, generates the meeting minutes, 
    # prints them, and then saves them into a Word document called meeting_minutes.docx.
audio_file_path = "Earningscall.wav"
transcription = transcribe_audio(audio_file_path)
minutes = meeting_minutes(transcription)
print(minutes)

save_as_docx(minutes, 'meeting_minutes.docx')